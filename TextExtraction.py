
from docx import Document
import pdfplumber

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            full_text.append(" | ".join(row_text))  
    
    return '\n'.join(full_text)

def extract_text_from_pdf(path):
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_txt(path):
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def extract_text_generic(path):
    raise ValueError("Unsupported file format: only .pdf and .docx are supported.")

def extract_text(path):
    if path.endswith(".pdf"):
        return extract_text_from_pdf(path)
    elif path.endswith(".docx"):
        return extract_text_from_docx(path)
    elif path.endswith(".txt"):
        return extract_text_from_txt(path)
    else:
        return extract_text_generic(path)

